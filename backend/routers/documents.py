import io
import requests
from pathlib import Path
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy.orm import Session
from typing import List, Optional
from datetime import datetime
from .. import models, schemas, database, auth
from ..cloudinary_utils import upload_file, delete_file

# Text extraction libraries
try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

router = APIRouter()


def extract_text_from_url(file_url: str, filename: str) -> str:
    """Download file from Cloudinary and extract text content."""
    ext = Path(filename).suffix.lower()
    
    try:
        resp = requests.get(file_url, timeout=30)
        resp.raise_for_status()
        file_bytes = resp.content
    except Exception as e:
        return f"Could not fetch file for preview: {str(e)}"
    
    try:
        if ext == '.txt':
            return file_bytes.decode('utf-8', errors='ignore')
        
        elif ext == '.pdf' and PdfReader:
            reader = PdfReader(io.BytesIO(file_bytes))
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            result = "\n\n".join(text_parts).strip()
            return result if result else "This PDF appears to be a scanned image or has no extractable text. Please download to view."
        
        elif ext == '.docx' and DocxDocument:
            doc = DocxDocument(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            result = "\n".join(paragraphs).strip()
            return result if result else "This document appears to be empty or contains only images. Please download to view."
        
        else:
            return f"Preview not available for {ext} files. Please download to view."
            
    except Exception as e:
        return f"Could not extract content: {str(e)}"


@router.post("/upload")
async def upload_document(
    title: str = Form(...),
    description: Optional[str] = Form(None),
    file_type: str = Form("general"),
    file: UploadFile = File(...),
    db: Session = Depends(database.get_db),
    admin: models.User = Depends(auth.require_admin)
):
    allowed_extensions = {'.pdf', '.doc', '.docx', '.xls', '.xlsx', '.csv', '.txt'}
    ext = Path(file.filename).suffix.lower()
    if ext not in allowed_extensions:
        raise HTTPException(status_code=400, detail=f"Invalid file type. Allowed: {allowed_extensions}")
    
    upload = upload_file(file, folder="lotsa/documents")
    
    # FIX: DB column is String, not Enum — pass the string directly
    db_doc = models.Document(
        title=title,
        description=description,
        file_type=file_type,
        file_url=upload["url"],
        file_public_id=upload["public_id"],
        file_name=file.filename,
        file_size=0,
        uploaded_by=admin.id
    )
    db.add(db_doc)
    db.commit()
    db.refresh(db_doc)
    return db_doc


@router.get("/", response_model=List[schemas.DocumentOut])
def list_documents(
    file_type: Optional[str] = None,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.get_current_active_user)
):
    query = db.query(models.Document).filter(models.Document.is_active == True)
    if file_type:
        query = query.filter(models.Document.file_type == file_type)
    return query.order_by(models.Document.uploaded_at.desc()).all()


@router.get("/admin/all", response_model=List[schemas.DocumentOut])
def admin_list_documents(
    file_type: Optional[str] = None,
    db: Session = Depends(database.get_db),
    admin: models.User = Depends(auth.require_admin)
):
    query = db.query(models.Document)
    if file_type:
        query = query.filter(models.Document.file_type == file_type)
    return query.order_by(models.Document.uploaded_at.desc()).all()


@router.get("/{doc_id}/content")
def get_document_content(
    doc_id: int,
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(auth.get_current_active_user)
):
    """Extract and return text content from an uploaded document."""
    doc = db.query(models.Document).filter(
        models.Document.id == doc_id,
        models.Document.is_active == True
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    text = extract_text_from_url(doc.file_url, doc.file_name)
    return {
        "id": doc.id,
        "title": doc.title,
        "file_name": doc.file_name,
        "content": text
    }


@router.delete("/{doc_id}")
def delete_document(
    doc_id: int,
    db: Session = Depends(database.get_db),
    admin: models.User = Depends(auth.require_admin)
):
    doc = db.query(models.Document).filter(models.Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if doc.file_public_id:
        delete_file(doc.file_public_id)
    
    db.delete(doc)
    db.commit()
    return {"message": "Document deleted"}


@router.put("/{doc_id}/toggle")
def toggle_document(
    doc_id: int,
    db: Session = Depends(database.get_db),
    admin: models.User = Depends(auth.require_admin)
):
    doc = db.query(models.Document).filter(models.Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    doc.is_active = not doc.is_active
    db.commit()
    db.refresh(doc)
    return doc